import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
import re
from scipy.stats import norm, ttest_ind
import io
from docx import Document
from docx.shared import Inches
import matplotlib.lines as mlines
import gc 

# =========================================================================
# 1. PAGE CONFIGURATION & FONTS
# =========================================================================
st.set_page_config(page_title="Line 4 Quality Analytics", layout="wide")

plt.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False

plt.rcParams.update({
    'font.size': 10,  
    'axes.labelweight': 'bold',
    'axes.titleweight': 'bold',
    'axes.titlesize': 13,
    'legend.fontsize': 9,
    'font.weight': 'bold',
    'lines.linewidth': 2.0,
    'figure.dpi': 100  
})

THEME_COLORS = ['#0055FF', '#FF6600', '#00AA00', '#9900FF', '#CC0055', '#009999']
MUTED_COLORS = ['#4A90E2', '#E67E22', '#27AE60', '#9B59B6', '#34495E', '#F1C40F']

# =========================================================================
# 2. UTILITY FUNCTIONS
# =========================================================================
@st.cache_data(max_entries=3) 
def load_and_clean_data(file):
    df = pd.read_csv(file) if file.name.endswith('.csv') else pd.read_excel(file)
    df.columns = [str(c).strip() for c in df.columns]
    return df

def find_data_col(df, key):
    for col in df.columns:
        if re.search(key, col, re.IGNORECASE) and not any(kw in col for kw in ["管制", "規格", "要求", "原始"]):
            return col
    return None

def get_limit(df, keyword, limit_type, category):
    col = next((c for c in df.columns if keyword in c and limit_type in c.lower() and category in c), None)
    if col:
        val = pd.to_numeric(df[col], errors='coerce').median()
        return float(val) if pd.notnull(val) and val > 0 else None
    return None

def get_limit_series(df, keyword, limit_type, category, length):
    col = next((c for c in df.columns if keyword in c and limit_type in c.lower() and category in c), None)
    if col:
        s = pd.to_numeric(df[col], errors='coerce')
        s = s.mask(s <= 0, np.nan).ffill().bfill()
        return s
    
    # [FIX LỖI INDEXING]: Gắn thêm index=df.index để đồng bộ với dataframe đã bị filter
    return pd.Series([np.nan] * length, index=df.index)

def format_num(val):
    if val is None or pd.isna(val): return "-"
    rounded = round(float(val), 2)
    return str(int(rounded)) if rounded == int(rounded) else str(rounded)

def format_spec(lsl, usl):
    if lsl != -1 and usl != -1:
        return f"{format_num(lsl)}-{format_num(usl)}"
    elif lsl != -1:
        return f"≥ {format_num(lsl)}"
    elif usl != -1:
        return f"≤ {format_num(usl)}"
    else:
        return "No Spec"

def apply_full_border(ax):
    for spine in ax.spines.values():
        spine.set_linewidth(2.0)
        spine.set_color('#111111')
        spine.set_visible(True)

def export_to_word(figures, titles):
    doc = Document()
    doc.add_heading('Quality Analytics Report', 0)

    for fig, title in zip(figures, titles):
        doc.add_heading(title, level=1)
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png', dpi=200, bbox_inches='tight')
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5.5))
        doc.add_paragraph("-" * 50)
    
    out_io = io.BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    return out_io

# HÀM MỚI: TÌM VÀ LÀM SẠCH DỮ LIỆU ĐỘ DÀY (THÔNG MINH HƠN)
def get_clean_thickness(df):
    thick_col = None
    
    # 1. Ưu tiên số 1: Tìm chính xác các cột độ dày chuẩn xác nhất
    priority_keywords = ["訂單厚度", "order thickness", "目標厚度", "target thickness"]
    for kw in priority_keywords:
        for c in df.columns:
            if kw.lower() in str(c).lower():
                thick_col = c
                break
        if thick_col: break
        
    # 2. Ưu tiên số 2: Nếu không tìm thấy cột chuẩn, tìm theo từ khóa rộng nhưng LOẠI TRỪ các cột nhiễu
    if not thick_col:
        exclude_keywords = ["公差", "套筒", "sleeve", "tolerance", "建議", "suggested"]
        for kw in ["厚度", "thickness", "thick", "độ dày"]:
            for c in df.columns:
                c_str = str(c).lower()
                # Kiểm tra có chứa từ khóa VÀ không chứa từ bị cấm
                if kw in c_str and not any(ex in c_str for ex in exclude_keywords):
                    thick_col = c
                    break
            if thick_col: break
    
    # 3. Trích xuất dữ liệu dạng số
    if thick_col:
        # Sử dụng RegEx để trích xuất chữ số (vd: "0.6mm" -> 0.6)
        num_series = pd.to_numeric(df[thick_col].astype(str).str.extract(r'(\d+\.?\d*)')[0], errors='coerce')
        return thick_col, num_series
        
    return None, pd.Series(dtype=float)

# =========================================================================
# 3. MAIN APP LOGIC & SIDEBAR
# =========================================================================
st.sidebar.header("📂 DATA SOURCE")
uploaded_files = st.sidebar.file_uploader("Upload Excel/CSV Reports", type=["xlsx", "csv", "xls"], accept_multiple_files=True)

if uploaded_files:
    metrics_map = {"YS": "YS", "TS": "TS", "EL": "EL", "Hardness": "HRB", "YPE": "YPE"}
    zh_map_global = {"YS": "降伏強度", "TS": "抗拉強度", "EL": "伸長率", "HRB": "硬度", "YPE": "YPE"}

    file_names = [f.name for f in uploaded_files]
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🏭 ASSIGN LINE DATA")
    ma_filename = st.sidebar.selectbox("Select Galvanizing File:", file_names, key="ma_sel")
    son_filename = st.sidebar.selectbox("Select Coating File:", file_names, index=1 if len(file_names) > 1 else 0, key="son_sel")
    
    st.sidebar.markdown("---")
    view_mode = st.sidebar.radio("🔍 VIEW MODE:", [
        "Process Analytics", 
        "SPC Control Charts (I-MR)", 
        "Executive Summary", 
        "Cross-Line Comparison 🔀"
    ])

    # =========================================================================
    # [VIEW MODE]: CROSS-LINE COMPARISON
    # =========================================================================
    if view_mode == "Cross-Line Comparison 🔀":
        st.title("🔀 Statistical Process Shift & Limits Recommendation")
        
        if ma_filename == son_filename:
            st.warning("⚠️ You selected the same file for both lines. Please assign different files in the sidebar to compare.")
        else:
            file_ma = next(f for f in uploaded_files if f.name == ma_filename)
            df_ma = load_and_clean_data(file_ma)
            
            file_son = next(f for f in uploaded_files if f.name == son_filename)
            df_son = load_and_clean_data(file_son)

            common_labels = [k for k, v in metrics_map.items() if find_data_col(df_ma, v) and find_data_col(df_son, v)]

            if not common_labels:
                st.error("❌ No common mechanical property columns found between the two files.")
            else:
                st.success(f"✅ Found {len(common_labels)} common properties to analyze: {', '.join(common_labels)}")
                
                for selected_label in common_labels:
                    st.markdown(f"<hr><h2 style='color: #2E86C1; text-align: center;'>🔹 Analysis for Parameter: {selected_label} 🔹</h2>", unsafe_allow_html=True)
                    
                    short_key = metrics_map[selected_label]
                    zh_key = zh_map_global.get(short_key, short_key)

                    col_ma = find_data_col(df_ma, short_key)
                    col_son = find_data_col(df_son, short_key)

                    thick_col_ma, thick_series_ma = get_clean_thickness(df_ma)
                    thick_col_son, thick_series_son = get_clean_thickness(df_son)
                    grade_col_ma = next((c for c in df_ma.columns if "等級" in str(c) or "grade" in str(c).lower()), None)
                    grade_col_son = next((c for c in df_son.columns if "等級" in str(c) or "grade" in str(c).lower()), None)

                    cols_ma = [col_ma] + ([thick_col_ma] if thick_col_ma else []) + ([grade_col_ma] if grade_col_ma else [])
                    temp_ma = df_ma[list(set(cols_ma))].copy()
                    temp_ma['val'] = pd.to_numeric(temp_ma[col_ma], errors='coerce')

                    if thick_col_ma:
                        temp_ma['Thick_Num'] = thick_series_ma
                        
                    if grade_col_ma:
                        temp_ma = temp_ma[temp_ma[grade_col_ma].astype(str).str.upper().str.contains('A|B|PRIME|1|2', na=True)]

                    temp_ma = temp_ma.dropna(subset=['val']).reset_index(drop=True)

                    cols_son = [col_son] + ([thick_col_son] if thick_col_son else []) + ([grade_col_son] if grade_col_son else [])
                    temp_son = df_son[list(set(cols_son))].copy()
                    temp_son['val'] = pd.to_numeric(temp_son[col_son], errors='coerce')

                    if thick_col_son:
                        temp_son['Thick_Num'] = thick_series_son

                    if grade_col_son:
                        temp_son = temp_son[temp_son[grade_col_son].astype(str).str.upper().str.contains('A|B|PRIME|1|2', na=True)]

                    temp_son = temp_son.dropna(subset=['val']).reset_index(drop=True)

                    groups_to_compare = []

                    if thick_col_ma and thick_col_son:
                        ma_g1 = temp_ma[temp_ma['Thick_Num'] <= 0.60].copy()
                        son_g1 = temp_son[temp_son['Thick_Num'] <= 0.60].copy()
                        if not ma_g1.empty and not son_g1.empty:
                            groups_to_compare.append(("Thickness <= 0.60", ma_g1, son_g1))
                            
                        ma_g2 = temp_ma[temp_ma['Thick_Num'] > 0.60].copy()
                        son_g2 = temp_son[temp_son['Thick_Num'] > 0.60].copy()
                        if not ma_g2.empty and not son_g2.empty:
                            groups_to_compare.append(("Thickness > 0.60", ma_g2, son_g2))
                    
                    if not groups_to_compare:
                        st.info(f"ℹ️ Configured thickness groups not found for {selected_label}. Switching to global analysis.")
                        groups_to_compare.append(("Global Data", temp_ma, temp_son))

                    for group_info in groups_to_compare:
                        group_name = group_info[0]
                        group_ma = group_info[1]
                        group_son = group_info[2]

                        st.markdown(f"<h3 style='color: #D35400;'>📌 Analysis: {group_name}</h3>", unsafe_allow_html=True)

                        vals_ma_full = group_ma['val']
                        vals_son_full = group_son['val']

                        mean_ma = vals_ma_full.mean()
                        mean_son = vals_son_full.mean()
                        
                        delta = mean_son - mean_ma if pd.notnull(mean_son) and pd.notnull(mean_ma) else 0

                        son_lsl_series = get_limit_series(df_son, zh_key, "min", "管制", len(df_son))
                        son_usl_series = get_limit_series(df_son, zh_key, "max", "管制", len(df_son))
                        
                        lsl_vals = son_lsl_series[son_lsl_series > 0]
                        usl_vals = son_usl_series[son_usl_series > 0]
                        
                        lsl_son = lsl_vals.mode()[0] if not lsl_vals.empty else None
                        usl_son = usl_vals.mode()[0] if not usl_vals.empty else None

                        if short_key == "YPE":
                            lsl_son = 4.0

                        s_lsl = (lsl_son - delta) if lsl_son is not None else "N/A"
                        s_usl = (usl_son - delta) if usl_son is not None else "N/A"

                        st.markdown(f"**🔄 Optimal Limits Recommendation**")
                        delta_data = [{
                            "Category": group_name,
                            "Galv. Theo. Value": format_num(mean_ma),
                            "Coating Theo. Value": format_num(mean_son),
                            "Shift (Δ)": format_num(delta),
                            "Current Coating LSL (Mode)": format_num(lsl_son) if lsl_son is not None else "N/A",
                            "Current Coating USL (Mode)": format_num(usl_son) if usl_son is not None else "N/A",
                            "Recommended Galv. LSL": format_num(s_lsl) if isinstance(s_lsl, (int, float)) else s_lsl,
                            "Recommended Galv. USL": format_num(s_usl) if isinstance(s_usl, (int, float)) else s_usl
                        }]
                        st.dataframe(pd.DataFrame(delta_data), hide_index=True, use_container_width=True)

                        if len(vals_son_full) > 1 and len(vals_ma_full) > 1:
                            t_stat, p_val = ttest_ind(vals_son_full, vals_ma_full, equal_var=False)
                            is_significant = "YES" if p_val < 0.05 else "NO"
                        else:
                            t_stat, p_val, is_significant = np.nan, np.nan, "N/A"
                        
                        c1, c2 = st.columns([1, 2])
                        with c1:
                            st.markdown("**🔬 2-Sample T-Test**")
                            t_test_data = pd.DataFrame([{
                                "Metric": "T-Statistic", "Value": format_num(t_stat)
                            }, {
                                "Metric": "P-Value", "Value": f"{p_val:.4f}" if pd.notnull(p_val) else "N/A"
                            }, {
                                "Metric": "Significant Shift?", "Value": is_significant
                            }])
                            st.table(t_test_data)

                        with c2:
                            fig_comp, ax_comp = plt.subplots(figsize=(8, 4))
                            
                            for label_name, vals, color in [
                                (f"Galvanizing (n={len(vals_ma_full)})", vals_ma_full, THEME_COLORS[0]),
                                (f"Coating (n={len(vals_son_full)})", vals_son_full, THEME_COLORS[1])
                            ]:
                                if len(vals) > 1 and vals.std() > 0:
                                    mu_val = vals.mean()
                                    sigma_val = vals.std(ddof=1)
                                    
                                    x_range = np.linspace(mu_val - 4*sigma_val, mu_val + 4*sigma_val, 500)
                                    bin_width = (vals.max() - vals.min()) / 20 if vals.max() > vals.min() else 1
                                    y_vals = norm.pdf(x_range, mu_val, sigma_val) * len(vals) * bin_width
                                    
                                    ax_comp.plot(x_range, y_vals, color=color, lw=2.5, label=label_name)
                                    ax_comp.fill_between(x_range, y_vals, alpha=0.3, color=color)
                                    ax_comp.axvline(mu_val, color=color, linestyle='--', alpha=0.8) 
                            
                            ax_comp.set_ylabel("Coil Count")
                            ax_comp.set_xlabel(f"{selected_label} Value")
                            
                            short_title = "Global" if lsl_son is None and usl_son is None else format_spec(lsl_son if lsl_son is not None else -1, usl_son if usl_son is not None else -1)
                            ax_comp.set_title(f"Shift Dist. (Δ = {format_num(delta)}) | {short_title}", pad=10)
                            
                            ax_comp.legend(loc="upper right", fontsize=9)
                            apply_full_border(ax_comp)
                            plt.tight_layout()
                            st.pyplot(fig_comp)
                            
                        plt.close(fig_comp)
                        gc.collect()
                        
                        st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================================
    # [VIEW MODE]: TABBED ANALYSIS SETUP (GLOBAL DATA PREP)
    # =========================================================================
    else:
        st.title(f"📊 {view_mode}")
        
        if view_mode == "SPC Control Charts (I-MR)":
            st.info("⚙️ Configure global target multipliers below. The system will apply them to all available properties across both lines.")
            c_i1, c_i2 = st.columns(2)
            with c_i1: k_std = st.number_input("Target Multiplier for StdDev (Sigma):", 1.0, 6.0, 3.0, 0.1)
            with c_i2: k_iqr = st.number_input("Target Multiplier for IQR (k-factor):", 1.0, 6.0, 1.5, 0.1)
        else:
            k_std, k_iqr = 3.0, 1.5

        tab_ma, tab_son = st.tabs(["🏭 Galvanizing Line Data", "🎨 Coating Line Data"])
        
        line_configs = [
            (tab_ma, ma_filename, "Galvanizing Line"),
            (tab_son, son_filename, "Coating Line")
        ]

        for tab_obj, fname, line_label in line_configs:
            with tab_obj:
                file_obj = next((f for f in uploaded_files if f.name == fname), None)
                if not file_obj:
                    st.warning(f"File '{fname}' not found.")
                    continue
                
                df_raw = load_and_clean_data(file_obj)
                
                necessary_cols = []
                for k, v in metrics_map.items():
                    col = find_data_col(df_raw, v)
                    if col: necessary_cols.append(col)
                for kw in ["管制", "規格", "要求", "厚度", "thickness", "用途碼", "thick", "độ dày"]:
                    cols = [c for c in df_raw.columns if kw in str(c).lower()]
                    necessary_cols.extend(cols)
                
                df = df_raw[list(set(necessary_cols))].copy() if necessary_cols else df_raw.copy()
                
                is_coating_line = any("原始" in str(c) for c in df_raw.columns)
                actual_line_type = "Coating Line" if is_coating_line else "Galvanizing Line"
                
                st.info(f"📂 Analyzing File: **{fname}** | Auto-detected: **{actual_line_type}**")
                
                time_col = next((c for c in df_raw.columns if "日期" in str(c) or "date" in str(c).lower() or "time" in str(c).lower()), None)
                date_range_str = "Unknown Period"
                if time_col:
                    temp_dates = pd.to_datetime(df_raw[time_col], errors='coerce').dropna()
                    if not temp_dates.empty:
                        date_range_str = f"{temp_dates.min().strftime('%Y-%m-%d')} to {temp_dates.max().strftime('%Y-%m-%d')}"
                
                st.markdown(f"**📅 Data Period:** `{date_range_str}`")
                st.markdown("---")
                
                if "用途碼" in df.columns:
                    usage_list = sorted(df["用途碼"].dropna().unique().tolist())
                    selected_usages = st.multiselect(f"Filter Usage Code ({line_label}):", options=usage_list, default=usage_list, key=f"usage_{fname}_{view_mode}")
                    df = df[df["用途碼"].isin(selected_usages)]

                # --- Thickness Range Filter Logic (Updated) ---
                if view_mode == "Process Analytics":
                    thick_col_filter, df['Thick_Num_Filter'] = get_clean_thickness(df)
                    
                    if thick_col_filter:
                        valid_thick = df['Thick_Num_Filter'].dropna()
                        
                        if not valid_thick.empty:
                            min_t = float(valid_thick.min())
                            max_t = float(valid_thick.max())
                            
                            if min_t < max_t:
                                c_f1, c_f2 = st.columns([1, 2]) 
                                with c_f1:
                                    selected_thick = st.slider(
                                        f"🎚️ Thickness Range ({line_label}) - {thick_col_filter}:", 
                                        min_value=min_t, 
                                        max_value=max_t, 
                                        value=(min_t, max_t), 
                                        step=0.01, 
                                        key=f"thick_slider_{fname}"
                                    )
                                df = df[(df['Thick_Num_Filter'] >= selected_thick[0]) & (df['Thick_Num_Filter'] <= selected_thick[1])]
                            else:
                                st.info(f"ℹ️ All coils in this dataset have the exact same thickness ({min_t}). Thickness filter is disabled.")
                        else:
                            st.warning(f"⚠️ Valid numerical data could not be parsed from column: '{thick_col_filter}'.")
                    else:
                        st.warning("⚠️ 'Thickness' column not found in this file (e.g. 訂單厚度). Thickness filter is disabled.")
                # ------------------------------------

                available = [k for k, v in metrics_map.items() if find_data_col(df, v)]
                if not available:
                    st.warning(f"⚠️ Mechanical property data column not found in this file.")
                    continue

                # =========================================================================
                # [VIEW MODE]: EXECUTIVE SUMMARY
                # =========================================================================
                if view_mode == "Executive Summary":
                    summary_data = []
                    for selected_label in available:
                        short_key = metrics_map[selected_label]
                        data_col = find_data_col(df, short_key)
                        zh_key = zh_map_global.get(short_key, short_key)
                        
                        if data_col:
                            temp_df = df[[data_col]].copy()
                            temp_df['val'] = pd.to_numeric(temp_df[data_col], errors='coerce')
                            temp_df = temp_df.dropna(subset=['val']).reset_index(drop=True)
                            
                            if temp_df.empty: continue
                            
                            thick_col, thick_series = get_clean_thickness(df)
                            
                            groups = []
                            if thick_col:
                                temp_df['Thick_Num'] = thick_series
                                g1 = temp_df[temp_df['Thick_Num'] <= 0.60]
                                g2 = temp_df[temp_df['Thick_Num'] > 0.60]
                                g_nan = temp_df[temp_df['Thick_Num'].isna()]
                                
                                if not g1.empty: groups.append(("<= 0.60", g1))
                                if not g2.empty: groups.append(("> 0.60", g2))
                                if not g_nan.empty: groups.append(("Unknown", g_nan))
                            else:
                                groups.append(("All", temp_df))

                            for g_name, group_df in groups:
                                p_data = group_df['val']
                                n_count = len(p_data)
                                if n_count == 0: continue
                                
                                mu_v = p_data.mean()
                                sig_v = p_data.std(ddof=1) if n_count > 1 else 0
                                
                                i_lsl_series = get_limit_series(df, zh_key, "min", "管制", len(df))
                                i_usl_series = get_limit_series(df, zh_key, "max", "管制", len(df))
                                
                                lsl_vals = i_lsl_series[i_lsl_series > 0]
                                usl_vals = i_usl_series[i_usl_series > 0]
                                
                                i_lsl = lsl_vals.mode()[0] if not lsl_vals.empty else None
                                i_usl = usl_vals.mode()[0] if not usl_vals.empty else None
                                
                                if is_coating_line and short_key == "YPE":
                                    i_lsl = 4.0
                                
                                cp, ca, cpk, formula, status = "-", "-", "-", "-", "N/A"
                                cpk_val = None
                                if pd.notnull(sig_v) and sig_v > 0:
                                    if i_usl is not None and i_lsl is not None:
                                        cp_v = (i_usl - i_lsl) / (6 * sig_v)
                                        cnt, half = (i_usl + i_lsl) / 2, (i_usl - i_lsl) / 2
                                        ca_v = (mu_v - cnt) / half
                                        cpk_val = cp_v * (1 - abs(ca_v))
                                        cp, ca, cpk, formula = format_num(cp_v), f"{ca_v*100:.1f}%", format_num(cpk_val), "Cp*(1-|Ca|)"
                                    elif i_usl is not None:
                                        cpk_val = (i_usl - mu_v) / (3 * sig_v); cpk, formula = format_num(cpk_val), "Cpu"
                                    elif i_lsl is not None:
                                        cpk_val = (mu_v - i_lsl) / (3 * sig_v); cpk, formula = format_num(cpk_val), "Cpl"
                                        
                                    if cpk_val is not None:
                                        if cpk_val < 1.0: status = "🔴 Action Required"
                                        elif 1.0 <= cpk_val < 1.33: status = "🟡 Acceptable"
                                        elif 1.33 <= cpk_val <= 2.0: status = "🟢 Excellent"
                                        else: status = "🔵 Over-engineered (>2.0)"
                                
                                summary_data.append({
                                    "Parameter": selected_label, 
                                    "Thickness": g_name,
                                    "N": n_count, "Theo. Value": format_num(mu_v), "StdDev (σ)": format_num(sig_v),
                                    "Int LSL": format_num(i_lsl), "Int USL": format_num(i_usl), 
                                    "Cp": cp, "Ca": ca, "Cpk": cpk, 
                                    "Cpk Formula": formula, "Status": status
                                })
                    
                    if summary_data:
                        st.dataframe(pd.DataFrame(summary_data), hide_index=True, use_container_width=True)

                # =========================================================================
                # [VIEW MODE]: PROCESS ANALYTICS & SPC CONTROL CHARTS
                # =========================================================================
                else:
                    for selected_label in available:
                        st.markdown(f"<hr><h3 style='color: #2E86C1;'>🔹 Parameter: {selected_label}</h3>", unsafe_allow_html=True)
                        short_key = metrics_map[selected_label]
                        data_col = find_data_col(df, short_key) 
                        zh_key = zh_map_global.get(short_key, short_key)
                        
                        if data_col:
                            temp_df = df[[data_col]].copy()
                            temp_df[data_col] = pd.to_numeric(temp_df[data_col], errors='coerce')
                            
                            plot_df = temp_df.dropna(subset=[data_col]).reset_index(drop=True)
                            plot_data = plot_df[data_col]
                            n = len(plot_data)

                            # ---------------------------------------------------------
                            # SUB-VIEW: PROCESS ANALYTICS (TREND & DIST)
                            # ---------------------------------------------------------
                            if view_mode == "Process Analytics":
                                int_lsl_series = get_limit_series(df, zh_key, "min", "管制", len(df)).loc[temp_df[data_col].notna()].reset_index(drop=True)
                                int_usl_series = get_limit_series(df, zh_key, "max", "管制", len(df)).loc[temp_df[data_col].notna()].reset_index(drop=True)
                                cust_lsl_series = get_limit_series(df, zh_key, "min", "客戶要求", len(df)).loc[temp_df[data_col].notna()].reset_index(drop=True)
                                cust_usl_series = get_limit_series(df, zh_key, "max", "客戶要求", len(df)).loc[temp_df[data_col].notna()].reset_index(drop=True)

                                if is_coating_line and short_key == "YPE":
                                    int_lsl_series = pd.Series([4.0] * n)
                                
                                temp_plot_df = plot_df.copy()
                                temp_plot_df['LSL_temp'] = int_lsl_series.fillna(-1).values
                                temp_plot_df['USL_temp'] = int_usl_series.fillna(-1).values
                                
                                groups = temp_plot_df.groupby(['LSL_temp', 'USL_temp'])
                                is_multi_group = len(groups) > 1
                                
                                calc_data = plot_data
                                mu = calc_data.mean()
                                sigma_fixed = calc_data.std(ddof=1)
                                safe_sigma = sigma_fixed if pd.notnull(sigma_fixed) and sigma_fixed > 0 else 1
                                safe_mu = mu if pd.notnull(mu) else plot_data.mean()

                                tab_trend, tab_dist = st.tabs([f"📈 {selected_label} Trend", f"📊 {selected_label} Distribution"])

                                with tab_trend:
                                    fig_t, ax_t = plt.subplots(figsize=(11, 5.5)) 
                                    x_coords = np.arange(1, n+1)

                                    if not int_lsl_series.isna().all() and not int_usl_series.isna().all():
                                        lower_bound = int_lsl_series.ffill().bfill()
                                        upper_bound = int_usl_series.ffill().bfill()
                                    else:
                                        lower_bound = pd.Series([-np.inf] * n)
                                        upper_bound = pd.Series([np.inf] * n)

                                    label_dict = {}
                                    def add_to_label(val, name, color):
                                        if pd.isna(val) or val <= 0: return
                                        val = round(val, 1) 
                                        if val not in label_dict: label_dict[val] = []
                                        if not any(item['name'] == name for item in label_dict[val]):
                                            label_dict[val].append({'name': name, 'color': color})

                                    if not cust_lsl_series.isna().all():
                                        for c_val in cust_lsl_series.dropna().unique():
                                            if c_val > 0: 
                                                ax_t.axhline(c_val, color="#00AA00", linestyle=":", linewidth=2.0, alpha=0.9)
                                                add_to_label(c_val, "Cust LSL", "#00AA00")
                                    if not cust_usl_series.isna().all():
                                        for c_val in cust_usl_series.dropna().unique():
                                            if c_val > 0: 
                                                ax_t.axhline(c_val, color="#00AA00", linestyle=":", linewidth=2.0, alpha=0.9)
                                                add_to_label(c_val, "Cust USL", "#00AA00")

                                    color_idx = 0
                                    for (lsl, usl), group in groups:
                                        c = THEME_COLORS[color_idx % len(THEME_COLORS)]
                                        
                                        c_mean = c if is_multi_group else "#0055FF"
                                        c_limit = c if is_multi_group else "#FF0000"
                                        
                                        mask = temp_plot_df.index.isin(group.index)
                                        spec_txt = format_spec(lsl, usl)
                                        
                                        group_mean = group[data_col].mean()
                                        
                                        ax_t.axhline(group_mean, color=c_mean, linestyle="-", linewidth=2.0, alpha=0.7, label="Group Theo. Value" if color_idx==0 else None)
                                        add_to_label(group_mean, "Theo. Value", c_mean)
                                        
                                        if lsl != -1: 
                                            ax_t.axhline(lsl, color=c_limit, linestyle="--", linewidth=2.0, alpha=1.0)
                                            add_to_label(lsl, "Int LSL", c_limit)
                                        if usl != -1: 
                                            ax_t.axhline(usl, color=c_limit, linestyle="--", linewidth=2.0, alpha=1.0)
                                            add_to_label(usl, "Int USL", c_limit)
                                            
                                        ax_t.scatter(x_coords[mask], plot_data[mask], color=c, s=40, edgecolor="black", linewidth=1.0, zorder=4, label=f"Data ({spec_txt})")
                                        color_idx += 1

                                    mask_out = (plot_data < lower_bound) | (plot_data > upper_bound)
                                    if mask_out.any():
                                        ax_t.scatter(x_coords[mask_out], plot_data[mask_out], color="#FF0000", s=60, edgecolor="#800000", linewidth=1.5, zorder=6, label="Out of Limit")

                                    valid_y = plot_data.dropna()
                                    ymin, ymax = valid_y.min(), valid_y.max()
                                    for val in label_dict.keys():
                                        ymin = min(ymin, val)
                                        ymax = max(ymax, val)
                                            
                                    y_range = ymax - ymin if ymax > ymin else 10
                                    ax_t.set_ylim(ymin - y_range*0.12, ymax + y_range*0.12)
                                    ax_t.set_xlim(0, n * 1.18)

                                    sorted_vals = sorted(label_dict.keys())
                                    min_y_dist = y_range * 0.05  
                                    last_y = -np.inf
                                    
                                    for val in sorted_vals:
                                        items = label_dict[val]
                                        names_str = " / ".join([item['name'] for item in items])
                                        main_color = items[0]['color']
                                        
                                        y_draw = val
                                        if y_draw - last_y < min_y_dist:
                                            y_draw = last_y + min_y_dist
                                            
                                        ax_t.plot([n, n + (n*0.02)], [val, y_draw], color="black", linestyle="-", lw=1.0, alpha=0.4)
                                        bbox = dict(boxstyle="round,pad=0.3", fc="#FFFFFF", ec=main_color, alpha=0.95, lw=1.5)
                                        ax_t.text(n + (n*0.025), y_draw, f"{names_str}: {val:.1f}", color=main_color, va='center', ha='left', fontsize=9, bbox=bbox, fontweight='bold')
                                        last_y = y_draw

                                    ax_t.set_xlabel("Coil Sequence")
                                    ax_t.set_ylabel(f"{selected_label} Value")
                                    ax_t.set_title(f"{selected_label} Trend Analysis (N={n})", pad=20)
                                    
                                    handles, labels = ax_t.get_legend_handles_labels()
                                    by_label = dict(zip(labels, handles))
                                    clean_dict = {k: v for k, v in by_label.items() if not k.startswith('_')}
                                    ax_t.legend(clean_dict.values(), clean_dict.keys(), loc="upper center", bbox_to_anchor=(0.5, -0.15), ncol=4, fontsize=9)
                                    
                                    apply_full_border(ax_t); plt.tight_layout(); st.pyplot(fig_t)
                                    
                                    buf_t = export_to_word([fig_t], [f"Trend Analysis - {selected_label}"])
                                    st.download_button(label=f"📥 Download Trend Chart ({selected_label})", data=buf_t, file_name=f"Trend_Report_{selected_label}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_trend_{fname}_{selected_label}")
                                    plt.close(fig_t)
                                
                                with tab_dist:
                                    # 1. Calculate safe bounds BEFORE plotting to avoid NameError
                                    x_min_data = plot_data.min()
                                    x_max_data = plot_data.max()
                                    x_range_data = x_max_data - x_min_data
                                    x_lower_bound = x_min_data - (x_range_data * 0.1)
                                    x_upper_bound = x_max_data + (x_range_data * 0.1)

                                    fig_d, ax_d = plt.subplots(figsize=(10, 6))
                                    plt.subplots_adjust(top=0.8, bottom=0.15) 
                                    
                                    # 2. Plot histograms
                                    hist_data, hist_labels = [], []
                                    for (lsl, usl), group in groups:
                                        hist_data.append(group[data_col].values)
                                        hist_labels.append(f"Data ({format_spec(lsl, usl)})")
                                        
                                    if is_multi_group:
                                        ax_d.hist(hist_data, bins=20, stacked=True, density=False, alpha=0.7, edgecolor="black", label=hist_labels, color=THEME_COLORS[:len(hist_data)])
                                    else:
                                        ax_d.hist(plot_data, bins=20, density=False, alpha=0.6, color="#0055FF", edgecolor="black", label="Data")

                                    ax_d.set_xlabel(f"{selected_label} Value")
                                    ax_d.set_ylabel("Coil Count")
                                    
                                    # 3. Twin axis for Normal Fit
                                    ax_pdf = ax_d.twinx()
                                    xs = np.linspace(x_lower_bound, x_upper_bound, 500)
                                    bin_w = (x_max_data - x_min_data) / 20 if x_range_data > 0 else 1
                                    ax_pdf.plot(xs, norm.pdf(xs, safe_mu, safe_sigma) * n * bin_w, color="#111111", lw=1.5, label="Normal Fit")
                                    ax_pdf.set_yticks([])
                                    
                                    # 4. Set ylim to leave space for labels
                                    ax_d.set_ylim(0, ax_d.get_ylim()[1] * 1.4) 
                                    
                                    # 5. Labeling logic
                                    lines_to_draw.sort(key=lambda x: x['val'])
                                    trans = ax_d.get_xaxis_transform()
                                    positions = [] 
                                    
                                    for item in lines_to_draw:
                                        val = item['val']
                                        ax_d.axvline(val, color=item['color'], linestyle=item['ls'], linewidth=2)
                                        
                                        y_top = ax_d.get_ylim()[1]
                                        y_pos = y_top * 0.85
                                        
                                        for prev_val, prev_y in positions:
                                            if abs(val - prev_val) < (x_range_data * 0.12):
                                                y_pos = prev_y + (y_top * 0.1)
                                        
                                        positions.append((val, y_pos))
                                        bbox = dict(boxstyle="round,pad=0.3", fc="white", ec=item['color'], alpha=0.9, lw=1)
                                        ax_d.text(val, y_pos, f"{val:.1f}", color=item['color'], ha='center', va='center', 
                                                  transform=trans, fontweight='bold', fontsize=9, bbox=bbox)
                                                  
                                    ax_d.set_title(f"{selected_label} Distribution (N={n})", pad=10, fontweight='bold')
                                    ax_d.legend(loc="upper left", bbox_to_anchor=(1.05, 1), fontsize=8)
                                    apply_full_border(ax_d)
                                    st.pyplot(fig_d)
                                    
                                    # 6. Export once
                                    buf_d = export_to_word([fig_d], [f"Distribution Analysis - {selected_label}"])
                                    st.download_button(label=f"📥 Download Dist Chart ({selected_label})", data=buf_d, file_name=f"Dist_Report_{selected_label}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_dist_{fname}_{selected_label}")
                                    plt.close(fig_d)
                                                                        
                                    buf_d = export_to_word([fig_d], [f"Distribution Analysis - {selected_label}"])
                                    st.download_button(label=f"📥 Download Dist Chart ({selected_label})", data=buf_d, file_name=f"Dist_Report_{selected_label}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_dist_{fname}_{selected_label}")
                                    plt.close(fig_d)
                                    
                            # ---------------------------------------------------------                            
                            # ---------------------------------------------------------
                            # ---------------------------------------------------------
                            # SUB-VIEW: SPC CONTROL CHARTS (I-MR)
                            # ---------------------------------------------------------
                            elif view_mode == "SPC Control Charts (I-MR)":
                                # 1. RE-FILTERING LOGIC (Ensure Slider is always active here)
                                thick_col, thick_series = get_clean_thickness(df)
                                temp_spc_df = df[[data_col]].copy()
                                temp_spc_df[data_col] = pd.to_numeric(temp_spc_df[data_col], errors='coerce')
                                
                                if thick_col:
                                    temp_spc_df['Thick_Num'] = thick_series
                                    min_t, max_t = float(temp_spc_df['Thick_Num'].min()), float(temp_spc_df['Thick_Num'].max())
                                    
                                    # This slider will always show when you are in SPC mode
                                    selected_thick = st.slider(
                                        f"🎚️ Thickness Range ({line_label}):", 
                                        min_value=min_t, max_value=max_t, value=(min_t, max_t), step=0.01,
                                        key=f"spc_slider_{fname}_{selected_label}"
                                    )
                                    temp_spc_df = temp_spc_df[
                                        (temp_spc_df['Thick_Num'].isna()) | 
                                        ((temp_spc_df['Thick_Num'] >= selected_thick[0]) & (temp_spc_df['Thick_Num'] <= selected_thick[1]))
                                    ]
                                
                                temp_spc_df = temp_spc_df.dropna(subset=[data_col]).reset_index(drop=True)
                                spc_groups = [("Filtered Data", temp_spc_df)] if not temp_spc_df.empty else []
                                    
                                st.markdown(f"#### 📐 Control Parameters Table")

                                # 2. DATA CALCULATION & TABLE
                                spc_stats = []
                                for g_name, group in spc_groups:
                                    g_data = group[data_col].dropna()
                                    if len(g_data) > 1:
                                        g_n, g_mu, g_sig = len(g_data), g_data.mean(), g_data.std(ddof=1)
                                        g_q1, g_q3 = g_data.quantile(0.25), g_data.quantile(0.75)
                                        g_iqr = g_q3 - g_q1
                                        spc_stats.append({
                                            "Category": g_name, "N": g_n, "Mean": format_num(g_mu), "Sigma": format_num(g_sig),
                                            f"Mill Range Upper ({k_std}σ)": format_num(g_mu + k_std*g_sig),
                                            f"Mill Range Lower ({k_std}σ)": format_num(g_mu - k_std*g_sig),
                                            "IQR": format_num(g_iqr),
                                            "UCL (IQR)": format_num(g_q3 + k_iqr*g_iqr),
                                            "LCL (IQR)": format_num(g_q1 - k_iqr*g_iqr)
                                        })
                                if spc_stats: st.dataframe(pd.DataFrame(spc_stats), hide_index=True, use_container_width=True)

                                # 3. CHART PLOTTING
                                fig_imr, ax_i = plt.subplots(figsize=(11, 5.5)) 
                                ax_i.plot(np.arange(1, len(temp_spc_df)+1), temp_spc_df[data_col], color="#CFD8DC", linestyle="-", linewidth=1.5, zorder=1)
                                
                                if spc_groups:
                                    g_mu = temp_spc_df[data_col].mean()
                                    g_sig = temp_spc_df[data_col].std(ddof=1)
                                    ax_i.scatter(np.arange(1, len(temp_spc_df)+1), temp_spc_df[data_col], color="#0055FF", s=40, edgecolor="black", zorder=3)
                                    ax_i.axhline(g_mu, color="#0055FF", linestyle="-", linewidth=2.0, label='Mean')
                                    ax_i.axhline(g_mu + k_std*g_sig, color="#FF0000", linestyle="--", linewidth=1.8, label=f'Mill Range ({k_std}σ)')
                                    ax_i.axhline(g_mu - k_std*g_sig, color="#FF0000", linestyle="--", linewidth=1.8)
                                
                                ax_i.set_xlabel("Coil Sequence"); ax_i.set_ylabel(f"{selected_label} Value")
                                ax_i.set_title(f"I-Chart: Dynamic Control Limits ({selected_label})", pad=20)
                                ax_i.legend(loc="upper left", bbox_to_anchor=(1, 1))
                                
                                apply_full_border(ax_i); plt.tight_layout(rect=[0, 0, 0.85, 1]); st.pyplot(fig_imr)
                                
                                buf_i = export_to_word([fig_imr], [f"SPC I-Chart Analysis - {selected_label}"])
                                st.download_button(label=f"📥 Download SPC Chart ({selected_label})", data=buf_i, file_name=f"SPC_Report_{selected_label}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_spc_{fname}_{selected_label}")
                                plt.close(fig_imr)
                                
                            gc.collect()
else:
    st.info("👈 Please upload the production report to start.")
